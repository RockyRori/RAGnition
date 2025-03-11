import os
from langchain_community.document_loaders import (
    PyPDFLoader,
    TextLoader,
    Docx2txtLoader,
)
from langchain_unstructured import UnstructuredLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from docx import Document as DocxDocument
from docx.shared import Pt
from fpdf import FPDF 
from markdownify import markdownify  

class UniversalDocumentSplitter:
    def __init__(
        self,
        input_folder,
        output_folder="splited",
        output_format="txt",  # 支持: txt/docx/md/html
        chunk_size=800,
        chunk_overlap=100
    ):
        self.input_folder = os.path.normpath(input_folder)
        self.output_folder = os.path.normpath(output_folder)
        self.output_format = output_format.lower()
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ".", " "]
        )
        os.makedirs(self.output_folder, exist_ok=True)

    def load_document(self, file_path):
        try:
            if file_path.endswith(".pdf"):
                return PyPDFLoader(file_path).load()
            elif file_path.endswith(".txt"):
                return TextLoader(file_path, encoding="utf-8").load()
            elif file_path.endswith(".docx"):
                return Docx2txtLoader(file_path).load()
            else:
                return UnstructuredLoader(file_path).load()
        except Exception as e:
            print(f"Error loading {file_path}: {str(e)}")
            return []

    def adaptive_split(self, document):
        paragraphs = [doc.page_content for doc in document]
        final_splits = []
        buffer = ""
        chunk_id = 1
        
        for p in paragraphs:
            buffer += p
            if len(buffer) > self.chunk_size:
                splits = self.splitter.split_text(buffer)
                for split in splits[:-1]:
                    final_splits.append({
                        "id": chunk_id,
                        "content": split,
                        "merged": False
                    })
                    chunk_id += 1
                buffer = splits[-1] if splits else ""
            elif len(buffer) < 300 and final_splits:
                # 合并逻辑
                final_splits[-1]["content"] += f"\n{buffer}"
                final_splits[-1]["merged"] = True
                buffer = ""
                chunk_id += 1  # 合并后递增编号
        if buffer:
            final_splits.append({
                "id": chunk_id,
                "content": buffer,
                "merged": False
            })
        return final_splits

    def save_document(self, file_path, splits):
        base_name = os.path.splitext(os.path.basename(file_path))[0]
        output_subdir = os.path.relpath(
            os.path.dirname(file_path),
            self.input_folder
        )
        output_dir = os.path.join(self.output_folder, output_subdir)
        os.makedirs(output_dir, exist_ok=True)
        
        # 根据格式选择保存方式
        if self.output_format == "docx":
            output_path = os.path.join(output_dir, f"{base_name}.docx")
            self.save_docx(splits, output_path)
        elif self.output_format == "md":
            output_path = os.path.join(output_dir, f"{base_name}.md")
            self.save_markdown(splits, output_path)
        elif self.output_format == "pdf":
            output_path = os.path.join(output_dir, f"{base_name}.pdf")
            self.save_pdf(splits, output_path)
        elif self.output_format == "html":
            output_path = os.path.join(output_dir, f"{base_name}.html")
            self.save_html(splits, output_path)
        else:
            # 默认保存为TXT
            output_path = os.path.join(output_dir, f"{base_name}.txt")
            self.save_txt(splits, output_path)

    def save_txt(self, splits, output_path):
        with open(output_path, "w", encoding="utf-8") as f:
            for chunk in splits:
                header = f"Chunk {chunk['id']}{' (merged)' if chunk['merged'] else ''}:\n"
                f.write(f"{header}{chunk['content']}\n\n")

    def save_docx(self, splits, output_path):
        doc = DocxDocument()
        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)
        
        for chunk in splits:
            # 添加标题
            title = f"Chunk {chunk['id']}{' (merged)' if chunk['merged'] else ''}:"
            doc.add_paragraph(title).bold = True
            # 添加内容
            doc.add_paragraph(chunk['content'])
            doc.add_paragraph()  # 分隔空行
        
        doc.save(output_path)

    def save_markdown(self, splits, output_path):
        with open(output_path, "w", encoding="utf-8") as f:
            for chunk in splits:
                header = f"### Chunk {chunk['id']}{' (merged)' if chunk['merged'] else ''}"
                content = f"```\n{chunk['content']}\n```"
                f.write(f"{header}\n{content}\n\n")

    def save_pdf(self, splits, output_path):
        pdf = FPDF()
        pdf.add_page()
        font_path = "DejaVuSans.ttf"
        if not os.path.exists(font_path):
            try:
                import requests
                from hashlib import md5

                url = "https://github.com/dejavu-fonts/dejavu-fonts/raw/master/ttf/DejaVuSans.ttf"
                expected_md5 = "a8c0b3b7b3d9a9e8f9d0e7f6c5b4a3d2"  # DejaVuSans.ttf 的正确 MD5

                response = requests.get(url)
                file_content = response.content
                actual_md5 = md5(file_content).hexdigest()

                if actual_md5 != expected_md5:
                    raise ValueError("字体文件损坏，请重试")

                with open(font_path, "wb") as f:
                    f.write(file_content)
            except Exception as e:
                raise RuntimeError(f"字体下载失败: {str(e)}，请手动下载 DejaVuSans.ttf")
        try:
            pdf.add_font('DejaVu', '', font_path, uni=True)
            pdf.set_font('DejaVu', size=12)
        except Exception as e:
            if "Not a TrueType font" in str(e):
                os.remove(font_path)  # 删除损坏的字体文件
                raise RuntimeError("检测到字体文件损坏，已自动删除，请重新运行程序")
            else:
                raise e
    def save_html(self, splits, output_path):
        html_content = """
        <html>
        <head><title>Document Chunks</title></head>
        <body style="font-family: Arial, sans-serif;">
        """
        
        for chunk in splits:
            style = "background-color: #f0f0f0; padding: 10px; margin: 10px 0;"
            title = f"Chunk {chunk['id']}{' (merged)' if chunk['merged'] else ''}:"
            content = chunk['content'].replace("\n", "<br>")
            html_content += f"""
            <div style="{style}">
                <h3>{title}</h3>
                <p>{content}</p>
            </div>
            """
        
        html_content += "</body></html>"
        
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(html_content)

    def process_folder(self):
        for root, _, files in os.walk(self.input_folder):
            for file in files:
                file_path = os.path.join(root, file)
                print(f"Processing {file_path}")
                document = self.load_document(file_path)
                if not document:
                    continue
                splits = self.adaptive_split(document)
                self.save_document(file_path, splits)

if __name__ == "__main__":
    splitter = UniversalDocumentSplitter(
        input_folder=r"D:\python_work\cantest",
        output_folder=r"D:\python_work\splited",
        output_format="txt",  # 支持: txt/docx/md/html
        chunk_size=600,
        chunk_overlap=100
    )
    splitter.process_folder()